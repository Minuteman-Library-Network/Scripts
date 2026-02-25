"""
Jeremy Goldstein with AI assistance from Claude.AI
Take bills notice job txt file from Sierra and divides into word docs by library
formatted for mailing, then uploads file to ftp server

Run on Demand
run in SIC
"""

import re
import os
import pysftp
import configparser
import time
from datetime import date
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def get_library_group_name(library_name):
    """
    Groups library branches together in order to produce consolidated files
    Returns the group name to use for filenames and the display name for documents.
    """
    library_upper = library_name.upper()

    if (
        "WEST ACTON BRANCH LIBRARY" in library_upper
        or "ACTON MEMORIAL LIBRARY" in library_upper
    ):
        return "ACT", "Acton"
    if "FOX BRANCH LIBRARY" in library_upper or "ROBBINS LIBRARY" in library_upper:
        return "ARL", "Arlington"
    if "ASHLAND PUBLIC LIBRARY" in library_upper:
        return "ASH", "Ashland"
    if "BEDFORD FREE PUBLIC LIBRARY" in library_upper:
        return "BED", "Bedford"
    if (
        "BELMONT LIBRARY AT BENTON LIB." in library_upper
        or "BELMONT LIBRARY AT BEECH ST." in library_upper
    ):
        return "BLM", "Belmont"
    if (
        "PUBLIC LIBRARY OF BROOKLINE" in library_upper
        or "COOLIDGE CORNER BRANCH LIBRARY" in library_upper
        or "PUTTERHAM BRANCH LIBRARY BROOKLINE" in library_upper
    ):
        return "BRK", "Brookline"
    if (
        "VALENTE BRANCH LIBRARY" in library_upper
        or "BOUDREAU BRANCH LIBRARY" in library_upper
        or "CAMBRIDGE PUBLIC LIBRARY-MAIN" in library_upper
        or "COLLINS BRANCH LIBRARY" in library_upper
        or "CONNELL BRANCH LIBRARY" in library_upper
        or "CENTRAL SQ. BRANCH LIBRARY" in library_upper
        or "NEILL BRANCH LIBRARY" in library_upper
    ):
        return "CAM", "Cambridge"
    if (
        "CONCORD PUBLIC LIBRARY" in library_upper
        or "LORING N FOWLER LIBRARY" in library_upper
    ):
        return "CON", "Concord"
    if "DEAN COLLEGE LIBRARY" in library_upper:
        return "DEA", "Dean"
    if (
        "DEDHAM ENDICOTT PUBLIC LIBRARY" in library_upper
        or "DEDHAM PUBLIC LIBRARY" in library_upper
    ):
        return "DDM", "Dedham"
    if "DOVER TOWN LIBRARY" in library_upper:
        return "DOV", "Dover"
    if (
        "FRAMINGHAM PUBLIC LIBRARY" in library_upper
        or "MCAULIFFE BRANCH LIBRARY" in library_upper
    ):
        return "FPL", "Framingham Public"
    if "HENRY WHITTEMORE LIBRARY" in library_upper:
        return "FST", "Framingham State"
    if "FRANKLIN PUBLIC LIBRARY" in library_upper:
        return "FRK", "Franklin"
    if "HOLLISTON PUBLIC LIBRARY" in library_upper:
        return "HOL", "Holliston"
    if "BRENNAN LIBRARY LASELL UNIVERSITY" in library_upper:
        return "LAS", "Lasell"
    if "CARY MEMORIAL LIBRARY" in library_upper:
        return "LEX", "Lexington"
    if "LINCOLN PUBLIC LIBRARY" in library_upper:
        return "LIN", "Lincoln"
    if "MAYNARD PUBLIC LIBRARY" in library_upper:
        return "MAY", "Maynard"
    if "MEDFIELD PUBLIC LIBRARY" in library_upper:
        return "MLD", "Medfield"
    if "MEDFORD PUBLIC LIBRARY" in library_upper:
        return "MED", "Medford"
    if "MEDWAY PUBLIC LIBRARY" in library_upper:
        return "MWY", "Medway"
    if "MILLIS PUBLIC LIBRARY" in library_upper:
        return "MIL", "Millis"
    if (
        "BACON FREE LIBRARY" in library_upper
        or "MORSE INSTITUTE LIBRARY" in library_upper
    ):
        return "NAT", "Natick"
    if "NEEDHAM FREE PUBLIC LIBRARY" in library_upper:
        return "NEE", "Needham"
    if "NEWTON FREE LIBRARY" in library_upper:
        return "NTN", "Newton"
    if "MORRILL MEMORIAL LIBRARY" in library_upper:
        return "NOR", "Norwood"
    if "OLIN COLLEGE LIBRARY" in library_upper:
        return "OLN", "Olin"
    if "REGIS COLLEGE LIBRARY" in library_upper:
        return "REG", "Regis"
    if "SHERBORN LIBRARY" in library_upper:
        return "SHR", "Sherborn"
    if (
        "EAST BRANCH LIBRARY" in library_upper
        or "SOMERVILLE PUBLIC LIBRARY" in library_upper
        or "WEST BRANCH LIBRARY" in library_upper
    ):
        return "SOM", "Somerville"
    if "RANDALL LIBRARY" in library_upper:
        return "STO", "Stow"
    if "GOODNOW PUBLIC LIBRARY" in library_upper:
        return "SUD", "Sudbury"
    if "WALTHAM PUBLIC LIBRARY" in library_upper:
        return "WLM", "Waltham"
    if "WATERTOWN FREE PUBLIC LIBRARY" in library_upper:
        return "WAT", "Watertown"
    if "WAYLAND PUBLIC LIBRARY" in library_upper:
        return "WYL", "Wayland"
    if (
        "WELLESLEY FREE LIBRARY" in library_upper
        or "HILLS BRANCH LIBRARY" in library_upper
        or "FELLS BRANCH LIBRARY" in library_upper
    ):
        return "WEL", "Wellesley"
    if "WESTON PUBLIC LIBRARY" in library_upper:
        return "WSN", "Weston"
    if (
        "WESTWOOD PUBLIC LIBRARY" in library_upper
        or "ISLINGTON BRANCH LIBRARY" in library_upper
    ):
        return "WWD", "Westwood"
    if "WINCHESTER PUBLIC LIBRARY" in library_upper:
        return "WIN", "Winchester"
    if "WOBURN PUBLIC LIBRARY" in library_upper:
        return "WOB", "Woburn"
    else:
        # For other libraries, use the original name
        clean_name = re.sub(r"[^\w\s-]", "", library_name)
        clean_name = re.sub(r"\s+", "_", clean_name)
        return clean_name, library_name


def create_mailing_ready_version(
    input_file_path, output_directory="C:\Scripts\\Bills\mailing_ready"
):
    """
    Creates mailing ready word doc of bills for a given library
    """

    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    try:
        with open(input_file_path, "r", encoding="utf-8") as file:
            content = file.read()
    except FileNotFoundError:
        print(f"Error: File '{input_file_path}' not found.")
        return

    library_pattern = r"([A-Z\s\'\.]+LIBRARY[A-Z\s\-\.]*)\n\s*(\d+[A-Z\s,\.\-\d]+)\n"
    matches = list(re.finditer(library_pattern, content))

    library_notifications = {}

    for i, match in enumerate(matches):
        start_pos = match.start()
        if i + 1 < len(matches):
            end_pos = matches[i + 1].start()
        else:
            end_pos = len(content)

        notification_text = content[start_pos:end_pos].rstrip()

        # Remove trailing page numbers to prevent separate pages
        lines = notification_text.split("\n")
        if lines and re.match(r"^\s*\d+:\d+\s*$", lines[-1]):
            notification_text = "\n".join(lines[:-1]).rstrip()

        library_name = match.group(1).strip()

        # Get the group name for this library
        group_name, display_name = get_library_group_name(library_name)

        if group_name not in library_notifications:
            library_notifications[group_name] = {
                "display_name": display_name,
                "notifications": [],
            }

        library_notifications[group_name]["notifications"].append(
            {
                "text": notification_text,
                "original_library": library_name,
                "index": i + 1,
            }
        )

    # Create mailing-ready Word documents
    for group_name, group_data in library_notifications.items():
        notifications = group_data["notifications"]

        filename = "{}Bills{}.docx".format(group_name, date.today().strftime("%b%d-%Y"))
        filepath = os.path.join(output_directory, filename)

        doc = Document()

        # Set precise margins for mailing
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        for idx, notification in enumerate(notifications):
            # Each notification gets its own page for individual mailing
            if idx > 0:
                doc.add_page_break()

            # Parse the notification and skip leading empty lines
            lines = notification["text"].split("\n")

            # Skip empty lines at the beginning of each notification
            start_idx = 0
            while start_idx < len(lines) and not lines[start_idx].strip():
                start_idx += 1

            # Process lines starting from first non-empty line
            for line_idx in range(start_idx, len(lines)):
                line = lines[line_idx]
                para = doc.add_paragraph()

                # Set paragraph formatting to preserve spacing
                para_format = para.paragraph_format
                para_format.space_before = Pt(0)
                para_format.space_after = Pt(0)
                para_format.line_spacing = 1.0

                # Add the line with Courier font
                run = para.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(10)

                # Bold important elements
                if any(keyword in line.upper() for keyword in ["LIBRARY"]):
                    run.bold = True
                elif "$" in line and "TOTAL" in line.upper():
                    run.bold = True

        try:
            doc.save(filepath)
            print(f"Mailing ready: {filename}")
        except Exception as e:
            print(f"Error saving mailing version {filename}: {e}")

        if "ACT" in group_name:
            display_name = "Acton"
        if "ARL" in group_name:
            display_name = "Arlington"
        if "ASH" in group_name:
            display_name = "Ashland"
        if "BED" in group_name:
            display_name = "Bedford"
        if "BLM" in group_name:
            display_name = "Belmont"
        if "BRK" in group_name:
            display_name = "Brookline"
        if "CAM" in group_name:
            display_name = "Cambridge"
        if "CON" in group_name:
            display_name = "Concord"
        if "DEA" in group_name:
            display_name = "Dean"
        if "DDM" in group_name:
            display_name = "Dedham"
        if "DOV" in group_name:
            display_name = "Dover"
        if "FPL" in group_name:
            display_name = "Framingham Public"
        if "FST" in group_name:
            display_name = "Framingham State"
        if "FRK" in group_name:
            display_name = "Franklin"
        if "HOL" in group_name:
            display_name = "Holliston"
        if "LAS" in group_name:
            display_name = "Lasell"
        if "LEX" in group_name:
            display_name = "Lexington"
        if "LIN" in group_name:
            display_name = "Lincoln"
        if "MAY" in group_name:
            display_name = "Maynard"
        if "MLD" in group_name:
            display_name = "Medfield"
        if "MED" in group_name:
            display_name = "Medford"
        if "MWY" in group_name:
            display_name = "Medway"
        if "MIL" in group_name:
            display_name = "Millis"
        if "NAT" in group_name:
            display_name = "Natick"
        if "NEE" in group_name:
            display_name = "Needham"
        if "NTN" in group_name:
            display_name = "Newton"
        if "NOR" in group_name:
            display_name = "Norwood"
        if "OLN" in group_name:
            display_name = "Olin"
        if "REG" in group_name:
            display_name = "Regis"
        if "SHR" in group_name:
            display_name = "Sherborn"
        if "SOM" in group_name:
            display_name = "Somerville"
        if "STO" in group_name:
            display_name = "Stow"
        if "SUD" in group_name:
            display_name = "Sudbury"
        if "WLM" in group_name:
            display_name = "Waltham"
        if "WAT" in group_name:
            display_name = "Watertown"
        if "WYL" in group_name:
            display_name = "Wayland"
        if "WELY" in group_name:
            display_name = "Wellesley"
        if "WSN" in group_name:
            display_name = "Weston"
        if "WWD" in group_name:
            display_name = "Westwood"
        if "WIN" in group_name:
            display_name = "Winchester"
        if "WOB" in group_name:
            display_name = "Woburn"
        ftp_file(filepath, display_name)


def preview_library_grouping(input_file_path):
    """
    Preview how notifications would be grouped by library, showing FPL grouping.
    """
    try:
        with open(input_file_path, "r", encoding="utf-8") as file:
            content = file.read()
    except FileNotFoundError:
        print(f"Error: File '{input_file_path}' not found.")
        return

    library_pattern = r"([A-Z\s\'\.]+LIBRARY[A-Z\s\-\.]*)\n\s*(\d+[A-Z\s,\.\-\d]+)\n"
    matches = list(re.finditer(library_pattern, content))

    library_count = {}
    group_count = {}

    for match in matches:
        library_name = match.group(1).strip()
        library_count[library_name] = library_count.get(library_name, 0) + 1

        # Get group information
        group_name, display_name = get_library_group_name(library_name)
        if group_name not in group_count:
            group_count[group_name] = {
                "display_name": display_name,
                "count": 0,
                "libraries": [],
            }
        group_count[group_name]["count"] += 1
        if library_name not in group_count[group_name]["libraries"]:
            group_count[group_name]["libraries"].append(library_name)

    print(
        f"Found {len(matches)} total notifications from {len(library_count)} individual libraries:"
    )
    print("-" * 60)

    for library, count in library_count.items():
        print(f"• {library}: {count} notification(s)")

    print(f"\nThese will be grouped into {len(group_count)} output files:")
    print("-" * 60)

    for group_name, group_info in group_count.items():
        print(f"• {group_info['display_name']}: {group_info['count']} notification(s)")
        if len(group_info["libraries"]) > 1:
            print(f"  Includes: {', '.join(group_info['libraries'])}")
        print(
            "  → Mailing:  {}Bills{}.docx".format(
                group_name, date.today().strftime("%b%d-%Y")
            )
        )
        print()


# upload report to SIC directory and optionally remove older files
def ftp_file(local_file, library, keep_local=False):

    config = configparser.ConfigParser()
    config.read("C:\\Scripts\\Creds\\config.ini")

    cnopts = pysftp.CnOpts()
    srv = pysftp.Connection(
        host=config["sic"]["sic_host"],
        username=config["sic"]["sic_user"],
        password=config["sic"]["sic_pw"],
        cnopts=cnopts,
    )

    local_file = local_file
    srv.cwd("/reports/Library-Specific Reports/" + library + "/Bills/")
    print("/reports/Library-Specific Reports/" + library + "/Bills/")
    srv.put(local_file)

    # remove old file

    retain = ["FY21", "FY22", "Fy22", "FY23", "fy23", "FY24", "Archive", "meta.json"]
    for fname in srv.listdir_attr():
        fullpath = (
            "/reports/Library-Specific Reports/"
            + library
            + "/Bills/{}".format(fname.filename)
        )
        # time tracked in seconds, st_mtime is time last modified
        name = str(fname.filename)
        if name not in retain and ((time.time() - fname.st_mtime) // (24 * 3600) >= 45):
            srv.remove(fullpath)

    srv.close()

    # Only remove local file if keep_local is False
    if not keep_local:
        os.remove(local_file)


# Example usage:
if __name__ == "__main__":
    # Replace with your actual file path
    input_file = "C:\Scripts\Bills\BillsNoEmail.txt"

    print("Preview of library grouping with FPL consolidation:")
    preview_library_grouping(input_file)

    print("\nCreating mailing-ready versions...")
    create_mailing_ready_version(input_file)
